from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "研究报告" / "基于HTML5的一种光学模拟实验平台_扩展版_无标签器材图.docx"


CN_ABSTRACT = (
    "摘要：传统光学实验教学通常依赖实体光具座、透镜、光屏和模拟眼等实验器材，"
    "在实际教学中容易受到设备数量、实验空间、调节精度、维护成本和课堂时间等因素限制。"
    "对于人眼屈光不正及其矫正等内容，学生往往只能观察有限的实验现象，难以连续、直观地理解"
    "焦点位置、视网膜成像、近视远视矫正和散光矫正之间的物理关系。与此同时，现代教育强调"
    "数字化教学资源建设、虚拟仿真实验应用和信息技术与实验教学深度融合，要求实验平台既能弥补"
    "实体实验条件不足，又能提升学生自主探究、数据分析和科学表达能力。基于上述背景，本项目设计并"
    "实现了一种基于 HTML5/WebGL 的模拟眼屈光不正及其矫正虚拟仿真实验平台。平台以几何光学和"
    "薄透镜模型为理论基础，结合 Three.js 三维可视化技术和 Blender 三维器材建模，构建了包含光源、"
    "物屏、双凸透镜、矫正镜片、模拟眼、像屏、光具座和光斑探测器的交互式实验场景。学生可选择不同"
    "屈光状态的模拟眼，调节像屏位置、镜片类型、镜片焦度和柱面镜角度，实时观察光路传播、焦点偏移、"
    "光斑清晰度和矫正效果。系统能够自动记录焦距测量数据，计算平均焦距、焦度和矫正镜片焦度，并生成"
    "可编辑、可打印的实验报告。该平台具有低成本、易部署、可重复操作、参数可视化、数据自动化和教学"
    "闭环完整等优点，可作为实体光学实验的有效补充，有助于提升学生对人眼成像与屈光矫正规律的理解。"
)

CN_KEYWORDS = "关键词：HTML5；WebGL；虚拟仿真实验；模拟眼；屈光不正；光学矫正；Three.js"

EN_ABSTRACT = (
    "Abstract: Traditional optics experiments usually rely on physical optical benches, lenses, screens and "
    "model-eye apparatus. In classroom practice, these experiments are often constrained by limited equipment, "
    "laboratory space, adjustment accuracy, maintenance cost and available teaching time. For topics such as "
    "refractive errors and optical correction of the human eye, students may observe only a limited range of "
    "phenomena and may find it difficult to build a continuous and visual understanding of focal position, retinal "
    "imaging, myopia and hyperopia correction, and astigmatism correction. In response to the educational trend "
    "toward digital teaching resources, virtual simulation experiments, and the integration of information "
    "technology with experimental teaching, this project develops an HTML5/WebGL-based virtual simulation platform "
    "for model-eye refractive errors and their correction. Based on geometrical optics and the thin-lens model, "
    "the platform combines Three.js visualization with Blender-modeled optical instruments to construct an "
    "interactive three-dimensional experimental scene, including a light source, object screen, convex lens, "
    "correction lenses, model eyes, image screen, optical bench and spot detector. Students can select different "
    "model eyes, adjust the screen position, lens type, lens power and cylindrical-lens angle, and observe ray "
    "propagation, focal shift, spot sharpness and correction effects in real time. The system records focal-length "
    "measurements automatically, calculates average focal length, optical power and correction power, and generates "
    "an editable and printable experiment report. The platform is low-cost, easy to deploy, repeatable, visually "
    "interactive and data-driven, providing an effective supplement to physical optics experiments and helping "
    "students better understand eye imaging and refractive correction."
)

EN_KEYWORDS = (
    "Keywords: HTML5; WebGL; virtual simulation experiment; model eye; refractive error; optical correction; Three.js"
)


def set_text(paragraph, text, bold_label=False):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    if bold_label:
        run.font.bold = True
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    set_text(inserted, text)
    return inserted


def main():
    doc = Document(DOCX)
    paragraphs = doc.paragraphs

    # Existing front matter: title, Chinese abstract placeholder, Chinese keywords, then blank paragraphs.
    set_text(paragraphs[2], CN_ABSTRACT)
    set_text(paragraphs[3], CN_KEYWORDS)

    # Clear old blank placeholders if present.
    for idx in [4, 5, 6]:
        if idx < len(paragraphs):
            set_text(paragraphs[idx], "")

    p = paragraphs[3]
    p = insert_after(p, EN_ABSTRACT)
    insert_after(p, EN_KEYWORDS)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()

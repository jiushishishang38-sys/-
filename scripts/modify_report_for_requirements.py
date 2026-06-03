from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\D\Desktop\26物理实验竞赛\仿真模拟")
SRC = ROOT / "研究报告" / "基于HTML5的模拟眼屈光不正及其矫正虚拟仿真实验平台 .docx"
OUT = ROOT / "研究报告" / "基于HTML5的模拟眼屈光不正及其矫正虚拟仿真实验平台_按考核要求修订版.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill="DCEBFF"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "宋体"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    r.font.size = Pt(10)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_i == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def insert_after(paragraph, new_block):
    paragraph._p.addnext(new_block._element)
    return new_block


def add_paragraph_after(doc, paragraph, text="", style=None):
    p = doc.add_paragraph(text, style=style)
    return insert_after(paragraph, p)


def add_table_after(doc, paragraph, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    paragraph._p.addnext(table._tbl)
    return table


def find_para(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise ValueError(f"Paragraph not found: {text}")


def add_requirements_table(doc):
    anchor = find_para(doc, "目标定位")
    p = add_paragraph_after(doc, anchor, "报告内容与考核要求对应关系", style="Heading 3")
    intro = add_paragraph_after(
        doc,
        p,
        "根据竞赛对设计报告的要求，本报告将以下七项内容分别落实到对应章节，便于评审快速核查作品的完整性。",
    )
    table = add_table_after(doc, intro, 8, 3)
    headers = ["要求项", "报告中的落实位置", "说明"]
    rows = [
        ["a 选题意义和目标定位", "题意解析及目标定位、引言、目标定位", "说明选题服务于人眼屈光实验教学，明确低成本、可交互、可推广的目标。"],
        ["b 相关物理原理", "实验原理", "覆盖几何光学、薄透镜公式、人眼模型、近视远视散光矫正和光线追迹。"],
        ["c 流程图与实现技术", "作品设计与实现", "补充程序总体流程图、模块结构、Three.js/WebGL/Blender/Vite 等实现技术。"],
        ["d 使用方法与参数范围", "使用方法", "说明模拟眼选择、像屏位置、镜片类型、焦度、柱面镜角度和数据记录方法。"],
        ["e 结果分析、合理性、有效性、可拓展性、局限与改进", "实验总结", "补充物理意义、有效性验证、拓展方向、局限和改进思路。"],
        ["f 电脑配置要求", "运行配置要求", "列出浏览器、显卡、Node.js、本地运行和静态部署要求。"],
        ["g 结论", "实验总结", "总结作品完成度、教学价值和后续优化方向。"],
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
    for r, row in enumerate(rows, start=1):
        for c, text in enumerate(row):
            set_cell_text(table.cell(r, c), text)
    style_table(table)


def add_implementation_flow(doc):
    anchor = find_para(doc, "3.2.3 作品开发流程")
    p = add_paragraph_after(doc, anchor, "3.2.3.1 程序总体流程图", style="Heading 3")
    flow = (
        "首页入口 -> 课前引导 -> 基础知识学习 -> 选择模拟眼与实验模式 -> "
        "读取参数 -> 计算焦点/焦度/光斑误差 -> 更新 Three.js 场景与光路 -> "
        "记录测量数据 -> 自动计算平均值和矫正焦度 -> 生成实验报告 -> 打印或保存 PDF"
    )
    flow_para = add_paragraph_after(doc, p, flow)
    for run in flow_para.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(38, 73, 129)
    note = add_paragraph_after(
        doc,
        flow_para,
        "该流程将教学环节、物理计算和界面反馈串联起来：学生每改变一次参数，系统都会重新评估成像位置和矫正状态，并把结果同步到三维光路、读数面板、数据表和报告草稿中。",
    )

    tech = add_paragraph_after(doc, note, "3.2.3.2 核心实现技术清单", style="Heading 3")
    table = add_table_after(doc, tech, 6, 3)
    rows = [
        ["技术环节", "采用技术", "作用"],
        ["三维显示", "Three.js、WebGL、OrbitControls、GLTFLoader", "搭建可旋转、可缩放的三维光具座与模拟眼实验场景。"],
        ["物理计算", "模块化 JavaScript、optics.js", "计算焦距、焦度、矫正镜片焦度、光斑误差和屈光类型。"],
        ["模型资源", "Blender、GLB 模型、PNG 人眼结构素材", "制作光具座、镜片支架、模拟眼和人眼结构展示资源。"],
        ["交互输入", "HTML5 表单、range 滑块、number 输入框、本地存储", "实现像屏位置、镜片焦度、柱面镜角度等参数调节和草稿保存。"],
        ["报告输出", "动态表格、localStorage、浏览器打印", "把实验数据同步到报告页面，并支持打印或保存为 PDF。"],
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.cell(r, c), text, bold=(r == 0))
    style_table(table)


def add_usage_parameters(doc):
    anchor = find_para(doc, "• 报告输出：在报告页补充分析与结论，使用打印功能保存为 PDF。")
    p = add_paragraph_after(doc, anchor, "主要参数设置范围与观察指标", style="Heading 2")
    table = add_table_after(doc, p, 7, 4)
    rows = [
        ["参数/控件", "建议范围或选项", "物理意义", "观察重点"],
        ["模拟眼编号", "A、B、C、D、E、F、G、S", "表示不同屈光状态；D 为正视眼，S 为散光模型。", "焦点相对 24.00 cm 视网膜位置的前后偏移。"],
        ["像屏位置", "约 18-32 cm 范围内细调", "模拟成像接收面位置。", "光斑最小时可认为接近清晰成像位置。"],
        ["矫正镜片类型", "无、凸透镜、凹透镜、柱面镜", "改变入射到模拟眼前的光线会聚或发散状态。", "近视用凹透镜，远视用凸透镜，散光用柱面镜。"],
        ["镜片焦度 D", "0-8 D，步进 0.25 D", "镜片改变光线传播方向的能力。", "比较计算焦度与实配焦度是否接近。"],
        ["柱面镜角度", "0-180°", "用于模拟散光矫正轴向。", "光斑椭圆程度和清晰方向随角度变化。"],
        ["记录次数", "每种模拟眼至少 3 次", "用于求平均值并减小偶然误差。", "平均焦距、焦度、屈光类型与误差分析。"],
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.cell(r, c), text, bold=(r == 0))
    style_table(table)


def add_result_analysis(doc):
    anchor = find_para(doc, "5.2 结论与展望")
    anchor.text = "5.4 结论"
    p = add_paragraph_after(doc, find_para(doc, "5.1 创新点"), "5.2 结果的物理意义与合理性分析", style="Heading 2")
    analysis_para = add_paragraph_after(
        doc,
        p,
        "仿真结果的核心物理意义在于把“焦点是否落在视网膜上”转化为可观察、可测量的实验量。当 A-C 型模拟眼的焦点位于 24.00 cm 视网膜位置之前时，系统判定为近视并推荐负焦度凹透镜；当 E-G 型模拟眼的焦点位于视网膜之后时，系统判定为远视并推荐正焦度凸透镜；D 型模拟眼焦点接近视网膜位置，表现为正视眼；S 型模拟眼则通过不同方向焦点不一致表现散光。该判据与几何光学和人眼屈光矫正原理一致，因此结果具有明确的物理合理性。",
    )

    p2 = add_paragraph_after(doc, analysis_para, "5.3 有效性、可拓展性、局限与改进", style="Heading 2")
    table = add_table_after(doc, p2, 5, 2)
    rows = [
        ["分析维度", "说明"],
        ["有效性", "平台能够实时显示焦点位置、光斑大小、矫正镜片类型和焦度计算值，学生可通过多次测量验证近视焦点前移、远视焦点后移和散光轴向相关的规律。"],
        ["可拓展性", "现有模块化结构便于继续加入真实镜片库、更多眼轴参数、教师端统计、实验成绩记录、多语言界面和移动端适配。"],
        ["局限性", "模型将真实人眼简化为等效薄透镜系统，未完整模拟角膜、晶状体、房水和玻璃体的多介质折射；光斑清晰度判断仍有教学近似；浏览器性能会影响三维渲染流畅度。"],
        ["改进思路", "后续可加入多界面折射模型、可调眼轴长度、真实视标图像成像、更多误差来源建模、数据导出和教师端批量评价功能。"],
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.cell(r, c), text, bold=(r == 0))
    style_table(table)


def add_runtime_table(doc):
    anchor = find_para(doc, "六、运行配置要求")
    p = add_paragraph_after(doc, anchor, "运行与部署配置表", style="Heading 2")
    table = add_table_after(doc, p, 6, 3)
    rows = [
        ["项目", "最低要求", "推荐配置/说明"],
        ["操作系统", "Windows 10 或同等级桌面系统", "Windows 11，或支持现代浏览器的 macOS/Linux。"],
        ["浏览器", "Microsoft Edge、Chrome、Firefox 新版本", "需开启 WebGL；建议使用 Chromium 内核浏览器以获得更稳定的 3D 性能。"],
        ["硬件", "4 GB 内存、普通集成显卡", "8 GB 以上内存，较新的集成显卡或独立显卡，便于流畅渲染 Three.js 场景。"],
        ["开发运行", "安装 Node.js 后执行 npm install、npm run dev", "本项目默认可通过 Vite 本地服务访问，地址可为 http://127.0.0.1:4173。"],
        ["静态部署", "执行 npm run build 后部署 dist 目录", "可部署到学校服务器、局域网电脑、GitHub Pages 或比赛展示电脑。"],
    ]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            set_cell_text(table.cell(r, c), text, bold=(r == 0))
    style_table(table)


def cleanup_blank_paragraphs(doc):
    for p in doc.paragraphs:
        if p.text.strip() == "":
            # Keep empty paragraphs but make them compact so insertions do not create visible gaps.
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    add_requirements_table(doc)
    add_implementation_flow(doc)
    add_usage_parameters(doc)
    add_result_analysis(doc)
    add_runtime_table(doc)
    cleanup_blank_paragraphs(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

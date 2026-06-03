from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "基于HTML5的一种光学模拟实验平台.docx"
OUT = ROOT / "研究报告" / "基于HTML5的一种光学模拟实验平台_扩展版_无标签器材图.docx"


def copy_table(dst, table):
    new_table = dst.add_table(rows=0, cols=len(table.columns))
    new_table.style = table.style
    for row in table.rows:
        new_row = new_table.add_row()
        for i, cell in enumerate(row.cells):
            new_cell = new_row.cells[i]
            new_cell._tc.get_or_add_tcPr()
            for p in list(new_cell.paragraphs):
                p._element.getparent().remove(p._element)
            for p in cell.paragraphs:
                new_p = new_cell.add_paragraph()
                new_p._p.getparent().remove(new_p._p)
                new_cell._tc.append(deepcopy(p._p))
    return new_table


def copy_paragraph(dst, paragraph):
    new_p = dst.add_paragraph()
    new_p._p.getparent().remove(new_p._p)
    dst._body._element.append(deepcopy(paragraph._p))
    return new_p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def style_doc(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    styles = doc.styles
    for name in ("Normal", "无缩进"):
        if name in styles:
            style = styles[name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
            style.paragraph_format.line_spacing = 1.25
            style.paragraph_format.space_after = Pt(6)
    for name, size, color in (
        ("Heading 1", 18, "000000"),
        ("Heading 2", 15, "000000"),
        ("Heading 3", 13, "000000"),
        ("Heading 4", 11, "000000"),
    ):
        if name in styles:
            style = styles[name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(6)


def add_p(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Inches(0.28) if style in (None, "Normal", "无缩进") else None
    p.add_run(text)
    return p


def add_bullet(doc, text):
    style = "List Bullet" if "List Bullet" in doc.styles else "List Paragraph" if "List Paragraph" in doc.styles else None
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.add_run("• " + text)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90, 90, 90)


def add_figure(doc, image_name, caption, width=5.9):
    path = ROOT / "screenshots" / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_code_block(doc, title, code):
    doc.add_paragraph(title, style="Heading 4")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(code.strip("\n").splitlines()):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8)


def add_module_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    headers = ["模块", "对应文件", "功能定位", "实现要点"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
    rows = [
        ("首页封面", "index.html", "平台入口与视觉导入", "视频封面、主导航、实验入口"),
        ("课前引导", "guide.html / guide.js", "实验背景和任务导入", "分步学习、动画过渡、学习目标呈现"),
        ("基础知识", "eye.html / eye-detailed.js", "人眼结构与屈光原理", "眼球结构图、部位说明、知识卡片"),
        ("模拟实验", "experiment.html / experiment.js", "核心三维仿真实验", "Three.js 场景、器材拖动、实时光路、光斑反馈"),
        ("光学模型", "optics.js", "屈光计算与数据处理", "焦度计算、屈光分类、推荐矫正、数据表渲染"),
        ("报告输出", "report.html / report.js", "实验报告填写与保存", "本地草稿保存、表格读取、打印导出"),
        ("模型资源", "model-assets.js / assets/models", "三维器材加载", "GLB 模型缓存、fallback 几何体、阴影与材质设置"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_margins(cells[i])


def add_blender_equipment_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    headers = ["3D 器材图片", "器材", "作用功能", "实现方法"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
    rows = [
        (
            "bench.png",
            "光具座与刻度尺",
            "作为实验器材的主支撑结构，提供物屏、透镜、模拟眼、像屏等元件沿光轴移动的空间参考。",
            "在 Blender 中建立长条形导轨、支撑底座和刻度面，导出 GLB 后在 Three.js 中叠加 Canvas 刻度纹理，并将厘米坐标映射到三维 x 轴。"
        ),
        (
            "source.png",
            "平行光源",
            "模拟实验中的入射光束，为观察透镜会聚、发散和人眼成像提供光线输入。",
            "建模时采用筒状灯体与发光面结构，网页端使用自发光材质和多条教学光线表示光束传播方向。"
        ),
        (
            "object.png",
            "物屏",
            "模拟被观察物体或靶标，为透镜成像与模拟眼成像提供物方对象。",
            "使用薄板与箭头/标记几何体组合建模，导出后绑定到可拖动支架，并通过位置参数控制物距。"
        ),
        (
            "convex.png",
            "双凸透镜",
            "用于准直与会聚光线，是构成实验光路和模拟眼前置光学系统的重要元件。",
            "在 Blender 中制作透明凸透镜外形，网页端设置玻璃材质、透明度、边缘高光，并在光路算法中按薄透镜近似更新光线斜率。"
        ),
        (
            "concave.png",
            "凹透镜",
            "作为近视矫正镜片，使入射光线适当发散，将焦点后移至视网膜附近。",
            "建模为中间薄、边缘厚的透明镜片，导入后与镜片支架联动显示；计算层将其焦度记为负值。"
        ),
        (
            "convex.png",
            "凸透镜矫正片",
            "作为远视矫正镜片，使入射光线预先会聚，将焦点前移至视网膜附近。",
            "与双凸透镜使用相近的玻璃模型，但在矫正支架中独立显示，焦度由滑块输入并参与矫正计算。"
        ),
        (
            "cylinder.png",
            "柱面镜",
            "用于散光矫正实验，表现不同轴向屈光能力不一致时的矫正过程。",
            "建立具有柱面方向特征的薄片模型，在 Three.js 中根据柱面镜角度旋转模型，并在光斑估算中引入轴向残余项。"
        ),
        (
            "support.png",
            "矫正镜片支架",
            "承载凹透镜、凸透镜或柱面镜，帮助学生观察矫正镜片加入前后的光路变化。",
            "在 Blender 中制作金属框架和底座，网页端根据镜片类型切换子模型可见性，保持支架位置固定在模拟眼前方。"
        ),
        (
            "eyeG.png",
            "A-G 模拟眼",
            "表示不同屈光能力的模拟眼：A-C 为近视型，D 为正视，E-G 为远视型。",
            "分别制作或配置不同编号的模拟眼 GLB 模型；程序根据 eyeId 加载对应模型，并通过 focusCm 参数参与屈光状态计算。"
        ),
        (
            "eyeD.png",
            "S 散光模拟眼",
            "用于展示散光状态下不同方向焦点不统一、光斑呈椭圆或弥散的现象。",
            "复用模拟眼结构并设置 astigmatic 标记，计算中结合柱面镜角度改变光斑椭圆率和视网膜清晰度判断。"
        ),
        (
            "screen.png",
            "像屏",
            "接收光线并显示成像清晰程度，用于寻找清晰焦点位置和记录焦距。",
            "建模为半透明薄板，绑定到光具座可拖动支架；网页端根据 screenCm 位置实时计算 screenError 和光斑大小。"
        ),
        (
            "screen.png",
            "光斑探测器",
            "用二维放大视图显示当前像屏上的光斑半径、椭圆程度和清晰度变化。",
            "三维场景中表现为探测器/显示面板，实际光斑由 HTML Canvas 绘制，使用径向渐变和椭圆参数模拟离焦与散光。"
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        image_name, *texts = row
        image_path = ROOT / "screenshots" / "latest-report" / "equipment" / image_name
        set_cell_margins(cells[0])
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if image_path.exists():
            cells[0].paragraphs[0].add_run().add_picture(str(image_path), width=Inches(1.3))
        else:
            cells[0].text = "图片待补充"
        for i, text in enumerate(texts, start=1):
            cells[i].text = text
            set_cell_margins(cells[i])


def append_expansion(doc):
    doc.add_page_break()
    doc.add_heading("三、作品的设计与实现", level=1)

    doc.add_heading("3.1 作品设计思路", level=2)
    add_p(doc, "本作品在原有透镜成像仿真的基础上，将实验主题进一步聚焦到“模拟眼屈光不正及其矫正”。平台并非只展示单一凸透镜或凹透镜成像，而是把光具座、准直透镜、矫正镜片、模拟眼和像屏组合成完整的人眼屈光实验系统，使学生能够在浏览器中完成从知识学习、参数调节、现象观察到数据记录和报告生成的完整实验过程。")
    add_p(doc, "作品设计遵循“知识铺垫—实验探究—数据分析—报告反思”的教学流程。学生先通过课前引导理解实验任务，再通过基础知识页面认识眼球结构和屈光原理，随后进入三维仿真实验台调节器材位置与镜片参数，最后在报告页面完成数据分析与结论表达。")
    add_bullet(doc, "可视化目标：将焦点前移、后移、光斑弥散、散光椭圆光斑等抽象现象转化为可观察画面。")
    add_bullet(doc, "交互目标：支持拖动光具座器材、输入精确数值、选择镜片类型和调节焦度。")
    add_bullet(doc, "教学目标：帮助学生理解近视、远视、散光的形成原因及矫正方法。")
    add_bullet(doc, "实验目标：实现焦距测量、平均值计算、焦度计算和矫正焦度比较。")

    add_figure(doc, "latest-report/01-index.png", "图3-1 平台首页与实验入口")
    add_figure(doc, "latest-report/02-guide.png", "图3-2 课前引导页面")

    doc.add_heading("3.2 作品实现技术", level=2)
    doc.add_heading("3.2.1 作品开发环境", level=3)
    add_p(doc, "作品运行环境为现代浏览器，开发环境为 Windows 本地开发环境。项目使用 Vite 构建前端工程，采用模块化 JavaScript 编写核心逻辑，页面由 HTML、CSS 与 JavaScript 共同实现。由于作品面向教学展示与比赛汇报，平台采用纯前端方式运行，便于部署到本地服务器或静态网站。")

    doc.add_heading("3.2.2 作品开发工具", level=3)
    add_p(doc, "平台开发主要使用 Visual Studio Code 进行代码编辑，使用 Three.js 完成 WebGL 三维场景搭建，使用 GLTFLoader 加载 GLB 三维模型资源，使用浏览器开发者工具进行交互调试与页面适配。部分三维器材模型通过 Blender 建模或处理后导出为 GLB 格式，用于增强实验场景的真实感。")

    doc.add_heading("3.2.3 作品开发流程", level=3)
    add_p(doc, "作品开发流程可分为需求分析、物理模型设计、页面结构搭建、三维场景实现、交互逻辑编写、数据记录与报告生成、测试与优化七个阶段。首先确定实验对象为模拟眼屈光不正及矫正；随后建立焦距、焦度、矫正镜片焦度之间的计算关系；再搭建多页面学习平台和三维实验台；最后通过自动化测试和页面截图检查交互稳定性。")

    doc.add_heading("3.2.4 前端页面与模块结构", level=3)
    add_module_table(doc)

    doc.add_heading("3.2.5 三维模型与场景搭建", level=3)
    add_p(doc, "三维实验台以光具座为主轴，按光传播方向依次布置光源、物屏、双凸透镜、矫正镜片支架、模拟眼光学组和像屏。场景中加入刻度尺、实验桌面、背景板和光斑显示器，使虚拟环境更接近真实物理实验室。用户可拖动可移动器材，也可通过右侧控制面板进行精确输入。")
    add_figure(doc, "latest-report/04-experiment.png", "图3-3 三维模拟实验台整体界面")

    doc.add_heading("3.2.6 Blender 建模光学器材介绍", level=3)
    add_p(doc, "为了使虚拟实验平台更接近真实光学实验环境，项目对主要光学器材进行了 Blender 建模，并以 GLB 格式导入网页端。建模时重点考虑器材的教学识别度、空间比例和交互绑定关系；在 Three.js 场景中，模型既承担视觉展示功能，也与光路计算、拖动控制和参数面板保持对应。")
    add_blender_equipment_table(doc)

    doc.add_heading("3.2.7 光学计算模型实现", level=3)
    add_p(doc, "平台将模拟眼等效为具有固定焦点位置的会聚光学系统，并设定视网膜位置为 24.00 cm。A、B、C 三种模拟眼焦点小于视网膜位置，表示屈光能力偏强，对应近视；D 为正视眼；E、F、G 焦点大于视网膜位置，表示屈光能力偏弱，对应远视；S 型模拟眼用于表现散光。系统根据焦距计算焦度，并根据模拟眼焦度与视网膜焦度的差值得出推荐矫正镜片焦度。")
    add_code_block(doc, "代码3-1 模拟眼参数与焦度计算", """
export const RETINA_CM = 24;

export const EYES = {
  A: { id: 'A', focusCm: 15.4, note: '高度屈光偏强' },
  B: { id: 'B', focusCm: 18.8, note: '中度屈光偏强' },
  C: { id: 'C', focusCm: 21.7, note: '轻度屈光偏强' },
  D: { id: 'D', focusCm: 24.0, note: '正视眼校准' },
  E: { id: 'E', focusCm: 27.6, note: '轻度屈光偏弱' },
  F: { id: 'F', focusCm: 31.4, note: '中度屈光偏弱' },
  G: { id: 'G', focusCm: 35.2, note: '高度屈光偏弱' },
  S: { id: 'S', focusCm: 24.0, astigmatic: true, note: '散光眼' }
};

export function diopterFromCm(cm) {
  return 100 / cm;
}

export function classifyEye(focusCm) {
  if (Math.abs(focusCm - RETINA_CM) < 0.55) return '正视眼';
  return focusCm < RETINA_CM ? '近视眼' : '远视眼';
}

export function correctionPowerForEye(focusCm) {
  return diopterFromCm(RETINA_CM) - diopterFromCm(focusCm);
}
""")

    add_code_block(doc, "代码3-2 实验状态评价与矫正判断", """
export function evaluateExperiment({ eyeId, screenCm, lensPower = 0, lensType = 'none', cylinderAngle = 0 }) {
  const eye = EYES[eyeId] || EYES.D;
  const signedLens = normalizeLensPower(lensType, lensPower);
  const effectivePower = diopterFromCm(eye.focusCm) + signedLens;
  const focusCm = 100 / Math.max(0.2, effectivePower);
  const cylinderActive = lensType === 'cylinder';
  const astigAngleTerm = Math.sin((cylinderAngle * Math.PI) / 90);
  const correctedAstigTerm = cylinderActive ? astigAngleTerm * 0.26 : astigAngleTerm;
  const target = eye.astigmatic ? RETINA_CM + correctedAstigTerm * 2.2 : RETINA_CM;
  const screenError = screenCm - focusCm;
  const retinaError = focusCm - target;
  return {
    eye,
    focusCm,
    screenError,
    retinaError,
    type: eye.astigmatic ? '散光眼' : classifyEye(eye.focusCm),
    correction: correctionPowerForEye(eye.focusCm),
    isClearOnScreen: Math.abs(screenError) < 0.8,
    isCorrected: Math.abs(retinaError) < 0.8
  };
}
""")

    doc.add_heading("3.3 仿真平台操作流程", level=2)
    add_p(doc, "仿真平台的操作流程按照真实实验步骤组织，既便于学生自主完成，也便于教师在课堂中进行演示。")
    for step in [
        "进入首页，选择课前引导、基础知识、模拟实验、实验报告或知识检测模块。",
        "在课前引导页面阅读实验目标，明确模拟眼屈光不正及矫正的探究任务。",
        "在基础知识页面认识眼球结构，理解角膜、晶状体和视网膜在人眼成像中的作用。",
        "进入模拟实验页面，选择 A-G 或 S 型模拟眼，调节像屏位置、物屏位置、准直透镜位置和矫正镜片参数。",
        "观察三维光路、光斑形态、当前焦点位置和系统给出的视网膜判断。",
        "点击记录按钮保存焦距测量值，完成三次测量后比较平均焦距、焦度和矫正值。",
        "进入实验报告页面填写现象观察、数据分析、结论反思，并打印或保存为 PDF。",
    ]:
        add_bullet(doc, step)
    add_figure(doc, "latest-report/03-eye.png", "图3-5 基础知识页面：眼球结构展示")
    add_figure(doc, "latest-report/05-report.png", "图3-6 实验报告页面")

    doc.add_heading("3.4 基础性实验：模拟眼屈光状态观察", level=2)
    add_p(doc, "基础性实验用于观察不同模拟眼未加矫正镜片时的成像状态。学生通过移动像屏寻找最清晰光斑位置，并将该位置作为模拟眼焦距测量值。若焦点小于 24.00 cm，说明焦点落在视网膜前，对应近视；若焦点大于 24.00 cm，说明焦点落在视网膜后，对应远视；若焦点接近 24.00 cm，则为正视眼。")
    add_p(doc, "该实验的价值在于让学生先从现象出发理解屈光状态差异，而不是直接记忆近视、远视的定义。通过 A-G 七种模拟眼的连续变化，学生可以看到屈光能力由强到弱时焦点位置的变化规律。")

    doc.add_heading("3.5 进阶性实验：屈光不正矫正仿真", level=2)
    add_p(doc, "进阶性实验主要研究矫正镜片对焦点位置的影响。对于近视型模拟眼，系统推荐凹透镜，使入射光线进入模拟眼前适当发散，从而使焦点后移至视网膜附近；对于远视型模拟眼，系统推荐凸透镜，使入射光线预先会聚，从而使焦点前移至视网膜附近。")
    add_p(doc, "散光实验则通过 S 型模拟眼和柱面镜角度调节实现。学生改变柱面镜角度时，可以观察到光斑椭圆程度和清晰度随轴向变化而改变，从而理解散光并不是单纯的焦点前后偏移，而是不同方向上的屈光能力不一致。")

    doc.add_heading("3.6 数据处理与分析", level=2)
    add_p(doc, "数据处理部分以焦距测量、平均值计算和焦度换算为核心。系统记录每个模拟眼三次测量值，并自动计算平均焦距、模拟眼焦度、屈光类型和矫正镜片焦度计算值。焦度计算公式为 D = 100 / f，其中 f 的单位为 cm，D 的单位为屈光度。")
    add_p(doc, "矫正镜片焦度近似由 D矫正 = D视网膜 - D模拟眼 得出。若结果为负，说明需要凹透镜矫正；若结果为正，说明需要凸透镜矫正。通过比较计算值与实际适配值，可以引导学生分析实验误差来源。")
    add_code_block(doc, "代码3-3 数据记录与表格计算", """
function applyRowCalculations(row, fittedPower = row.correctionFit) {
  if (!row.measurements.length) {
    resetRow(row);
    return row;
  }
  const average = row.measurements.reduce((sum, item) => sum + item, 0) / row.measurements.length;
  const diopter = diopterFromCm(average);
  const correction = correctionPowerForEye(average);
  row.average = average.toFixed(2);
  row.diopter = diopter.toFixed(2);
  row.type = classifyEye(average);
  row.correctionCalc = correction.toFixed(2);
  const fitted = Number(fittedPower);
  row.correctionFit = Number.isFinite(fitted) ? fitted.toFixed(2) : '';
  return row;
}
""")

    doc.add_heading("3.7 作品特色与创新点", level=2)
    add_bullet(doc, "将传统光具座实验、模拟眼实验和屈光矫正实验整合到同一个 Web 仿真平台中。")
    add_bullet(doc, "使用 Three.js 构建三维实验环境，使光路传播、焦点移动和光斑变化直观可见。")
    add_bullet(doc, "实现参数调节与物理计算实时联动，学生改变镜片或像屏位置后即可得到反馈。")
    add_bullet(doc, "支持实验数据自动记录、平均值计算、焦度换算和报告生成，形成完整教学闭环。")
    add_bullet(doc, "平台可通过浏览器运行，部署成本低，适合课堂演示、学生自学和实验竞赛展示。")

    doc.add_heading("四、使用方法", level=1)
    add_p(doc, "本平台既可以通过本地开发服务器运行，也可以在构建后作为静态网页部署。使用时先进入首页，再按照“课前引导—基础知识—模拟实验—实验报告—知识检测”的顺序完成学习。")
    add_bullet(doc, "本地运行：在项目目录中执行 npm run dev，浏览器访问 http://127.0.0.1:4173。")
    add_bullet(doc, "实验操作：选择模拟眼编号，调节像屏位置和矫正镜片参数，观察读数面板和光斑变化。")
    add_bullet(doc, "数据记录：点击“记录本次焦距”，完成多次测量后进入报告页查看结果。")
    add_bullet(doc, "报告输出：在报告页补充分析与结论，使用打印功能保存为 PDF。")

    doc.add_heading("五、实验总结", level=1)
    doc.add_heading("5.1 创新点", level=2)
    add_p(doc, "作品以 HTML5/WebGL 为技术基础，在浏览器中实现三维光学仿真，避免了传统虚拟实验对专用软件环境的依赖。与普通二维光路图相比，本平台能够展示器材空间位置、光具座刻度、实时光路和光斑清晰度，使学生更容易建立实验装置与物理模型之间的联系。")
    doc.add_heading("5.2 结论与展望", level=2)
    add_p(doc, "本作品完成了模拟眼屈光不正及其矫正实验的主要功能，实现了不同模拟眼成像状态观察、凹透镜和凸透镜矫正、柱面镜散光矫正、实验数据记录和报告生成。后续可进一步加入更精细的人眼模型、真实镜片参数库、教师端数据统计和多人实验记录功能，使平台更适合规模化教学应用。")

    doc.add_heading("六、运行配置要求", level=1)
    add_bullet(doc, "硬件要求：普通 Windows 电脑即可运行，建议具备独立或较新的集成显卡以获得更流畅的三维渲染效果。")
    add_bullet(doc, "软件要求：现代浏览器，如 Microsoft Edge、Chrome 或 Firefox。")
    add_bullet(doc, "开发运行：需安装 Node.js，并在项目目录中执行 npm install 与 npm run dev。")
    add_bullet(doc, "部署方式：可通过 Vite 构建为静态资源后部署到学校网站、GitHub Pages 或本地服务器。")

    doc.add_heading("七、参考文献", level=1)
    for ref in [
        "[1] 普通高中物理课程标准及相关教材中关于几何光学、人眼成像与视力矫正的内容。",
        "[2] Three.js 官方文档：WebGL 三维场景、材质、光照与模型加载相关说明。",
        "[3] Vite 官方文档：前端工程构建与静态资源部署说明。",
        "[4] 高等教育出版社新形态教材网相关虚拟仿真实验项目展示页面。",
    ]:
        add_p(doc, ref)

    doc.add_heading("八、附录", level=1)
    add_p(doc, "小组分工：可在此处补充项目策划、物理建模、前端开发、三维建模、测试与报告撰写等成员分工。")
    add_p(doc, "成本核算：本项目主要基于开源前端技术和已有浏览器环境开发，硬件成本较低，适合在普通机房或学生个人电脑中运行。")
    add_p(doc, "实验装置：虚拟装置包括光源、物屏、双凸透镜、矫正镜片支架、模拟眼、像屏、光具座和光斑探测面板。")


def main():
    dst = Document(SRC)
    style_doc(dst)

    replacements = {
        "研究几何光学中透镜成像的物理规律；利用HTML5/WebGL技术开发一个交互式光学模拟实验平台，实现凸透镜、凹透镜成像过程的可视化与参数化探究。":
            "研究几何光学中人眼成像与屈光矫正的物理规律；利用 HTML5/WebGL 技术开发一个交互式光学模拟实验平台，实现模拟眼屈光状态观察、近视远视矫正、散光矫正、实验数据记录与报告生成。",
        "本作品基于HTML5技术栈，构建3D可视化场景，自主编写光线追迹核心算法，实现凸透镜与凹透镜成像规律的动态模拟。平台支持用户拖拽调节物距、透镜焦距等参数，实时观察成像变化，并提供理想透镜与实际透镜（球差）两种模式对比。同时，系统自动记录实验数据、计算不确定度、生成可打印实验报告，辅助学生完成从探究学习到实验报告撰写的完整教学过程。":
            "本作品基于 HTML5 技术栈，构建 3D 可视化实验场景，自主设计模拟眼屈光计算与教学光路追踪逻辑，实现正视眼、近视眼、远视眼和散光眼的成像状态观察与矫正仿真。平台支持用户拖拽调节光具座器材、选择矫正镜片类型、调节镜片焦度和柱面镜角度，实时观察光路、焦点位置与光斑清晰度变化。同时，系统自动记录实验数据、计算平均焦距和矫正焦度，并生成可打印实验报告，辅助学生完成从探究学习到实验报告撰写的完整教学过程。",
        "本实验的主要目标为：以几何光学中的透镜成像规律为核心教学内容，利用HTML5/WebGL技术开发一个交互式光学模拟实验平台，实现3D可视化光线追迹、多模式成像对比、实验数据自动记录与报告生成，为大学物理实验教学提供低成本、可推广的虚拟仿真工具。":
            "本实验的主要目标为：以几何光学中的人眼屈光成像和视力矫正规律为核心教学内容，利用 HTML5/WebGL 技术开发一个交互式光学模拟实验平台，实现 3D 可视化光路演示、模拟眼屈光状态判断、矫正镜片参数探究、实验数据自动记录与报告生成，为物理实验教学提供低成本、可推广、可反复操作的虚拟仿真工具。"
    }
    for paragraph in dst.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(replacements[text])

    stop_headings = {"技术路线及功能实现", "使用方法", "实验总结", "运行配置要求", "参考文献", "附录"}
    body = dst._body._element
    children = list(body)
    remove_from = None
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(node.text or "" for node in child.iter() if node.tag == qn("w:t")).strip()
        if text in stop_headings:
            remove_from = idx
            break
    if remove_from is not None:
        for child in children[remove_from:]:
            if child.tag == qn("w:sectPr"):
                continue
            body.remove(child)

    append_expansion(dst)
    OUT.parent.mkdir(exist_ok=True)
    dst.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

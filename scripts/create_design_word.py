from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "研究报告" / "作品设计与实现_代码与截图.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_code_block(doc, title, code):
    p = doc.add_paragraph()
    p.style = "Heading 3"
    p.add_run(title)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "DADCE0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, 120, 160, 120, 160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            para.add_run("\n")
        run = para.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(35, 35, 35)

    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_figure(doc, image_name, caption, width=6.1):
    path = ROOT / "screenshots" / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in (
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("作品的设计与实现")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("眼屈光不正及其矫正仿真平台：关键设计、实现代码与页面截图")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)


def main():
    doc = Document()
    configure_doc(doc)
    add_title(doc)

    doc.add_heading("3.1 作品设计思路", level=1)
    doc.add_paragraph(
        "本作品以“模拟眼屈光不正及其矫正”为核心实验内容，面向中学物理光学实验教学中"
        "人眼成像、近视远视矫正、散光矫正等知识点，设计了集知识学习、三维仿真、交互操作、"
        "数据记录和实验报告生成于一体的数字化实验平台。"
    )
    add_bullet(doc, "课前引导：明确实验目标、学习任务和操作路径。")
    add_bullet(doc, "基础知识：展示角膜、虹膜、晶状体、玻璃体、视网膜、视神经等眼球结构。")
    add_bullet(doc, "模拟实验：通过三维光具座完成模拟眼选择、像屏调节、镜片矫正和光斑观察。")
    add_bullet(doc, "数据报告：自动记录焦距、计算焦度与矫正值，并形成可打印实验报告。")

    add_figure(doc, "cover-redesign-desktop.png", "图 3-1 平台首页与实验入口设计")
    add_figure(doc, "guide-react-gsap-4173-final.png", "图 3-2 课前引导页面：实验背景与任务导入")

    doc.add_heading("3.2 作品实现技术", level=1)
    doc.add_heading("3.2.1 开发环境与开发工具", level=2)
    doc.add_paragraph(
        "项目采用网页端实现方式，运行于现代浏览器。前端构建工具为 Vite，主要开发语言为 "
        "JavaScript、HTML 与 CSS，三维场景渲染使用 Three.js。模型资源采用 GLB 格式，"
        "通过 GLTFLoader 加载到网页场景中。"
    )
    add_bullet(doc, "开发环境：Windows、本地浏览器、Node.js/Vite。")
    add_bullet(doc, "核心库：Three.js、OrbitControls、GLTFLoader。")
    add_bullet(doc, "资源类型：HTML 页面、CSS 样式、JavaScript 模块、GLB 三维模型、PNG 页面截图。")

    doc.add_heading("3.2.2 页面与模块结构", level=2)
    module_table = doc.add_table(rows=1, cols=3)
    module_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(module_table)
    headers = ["模块", "对应页面/脚本", "主要功能"]
    for i, text in enumerate(headers):
        cell = module_table.cell(0, i)
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.text = text
    rows = [
        ("首页封面", "index.html", "平台入口、视觉导入与导航。"),
        ("课前引导", "guide.html / guide.js", "实验背景、任务流程与学习准备。"),
        ("基础知识", "eye.html / eye-detailed.js", "眼球结构与屈光原理展示。"),
        ("模拟实验", "experiment.html / experiment.js", "三维光具座、光路更新、数据记录。"),
        ("实验报告", "report.html / report.js", "实验数据表、草稿保存、打印导出。"),
        ("知识检测", "quiz.html / quiz.js", "实验后知识巩固与学习反馈。"),
    ]
    for row in rows:
        cells = module_table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_margins(cells[i])

    doc.add_heading("3.2.3 光学计算模型实现", level=2)
    doc.add_paragraph(
        "平台将人眼近似为会聚成像系统，设定视网膜位置为 24.00 cm。A-G 七种模拟眼通过不同焦点位置"
        "表示不同屈光状态，系统根据焦距计算焦度，并进一步给出推荐矫正镜片。"
    )
    add_code_block(
        doc,
        "代码片段 1：模拟眼参数、焦度计算与屈光状态判断",
        """
export const RETINA_CM = 24;

export const EYES = {
  A: { id: 'A', focusCm: 15.4, note: '高度屈光偏强' },
  B: { id: 'B', focusCm: 18.8, note: '中度屈光偏强' },
  C: { id: 'C', focusCm: 21.7, note: '轻度屈光偏强' },
  D: { id: 'D', focusCm: 24.0, note: '正视眼校准' },
  E: { id: 'E', focusCm: 27.6, note: '轻度屈光偏弱' },
  F: { id: 'F', focusCm: 31.4, note: '中度屈光偏弱' },
  G: { id: 'G', focusCm: 35.2, note: '高度屈光偏弱' },
  S: { id: 'S', focusCm: 24.0, astigmatic: true, note: '散光眼' }
};

export function diopterFromCm(cm) {
  return 100 / cm;
}

export function classifyEye(focusCm) {
  if (Math.abs(focusCm - RETINA_CM) < 0.55) return '正视眼';
  return focusCm < RETINA_CM ? '近视眼' : '远视眼';
}

export function correctionPowerForEye(focusCm) {
  return diopterFromCm(RETINA_CM) - diopterFromCm(focusCm);
}

export function recommendedLens(power) {
  if (Math.abs(power) < 0.15) return { type: 'none', label: '无需矫正' };
  return power < 0
    ? { type: 'concave', label: '凹透镜' }
    : { type: 'convex', label: '凸透镜' };
}
""",
    )

    add_code_block(
        doc,
        "代码片段 2：实验状态评价与光斑清晰度估算",
        """
export function evaluateExperiment({
  eyeId,
  screenCm,
  lensPower = 0,
  lensType = 'none',
  cylinderAngle = 0
}) {
  const eye = EYES[eyeId] || EYES.D;
  const signedLens = normalizeLensPower(lensType, lensPower);
  const effectivePower = diopterFromCm(eye.focusCm) + signedLens;
  const focusCm = 100 / Math.max(0.2, effectivePower);

  const cylinderActive = lensType === 'cylinder';
  const astigAngleTerm = Math.sin((cylinderAngle * Math.PI) / 90);
  const correctedAstigTerm = cylinderActive ? astigAngleTerm * 0.26 : astigAngleTerm;
  const target = eye.astigmatic ? RETINA_CM + correctedAstigTerm * 2.2 : RETINA_CM;

  const screenError = screenCm - focusCm;
  const retinaError = focusCm - target;
  const spotRadius = estimateSpotRadiusCm(screenError, eye.astigmatic ? Math.abs(correctedAstigTerm) : 0);

  return {
    eye,
    focusCm,
    screenError,
    retinaError,
    spotRadius,
    type: eye.astigmatic ? '散光眼' : classifyEye(eye.focusCm),
    correction: correctionPowerForEye(eye.focusCm),
    recommended: recommendedLens(correctionPowerForEye(eye.focusCm)),
    isClearOnScreen: Math.abs(screenError) < 0.8,
    isCorrected: Math.abs(retinaError) < 0.8
  };
}
""",
    )

    doc.add_heading("3.2.4 三维模型与场景搭建", level=2)
    doc.add_paragraph(
        "三维实验台由光具座、光源、物屏、双凸透镜、矫正镜片支架、模拟眼、像屏和光斑显示面板组成。"
        "各器材沿光具座方向排列，用户可以通过滑块、数值框或鼠标拖动改变器材位置。"
    )
    add_figure(doc, "verify-experiment.png", "图 3-3 三维模拟实验平台整体界面")
    add_figure(doc, "preset-detector-spot.png", "图 3-4 光斑显示与清晰度反馈")
    add_code_block(
        doc,
        "代码片段 3：三维实验状态读取与实时更新",
        """
function readExperimentState() {
  return {
    mode: modeInput.value,
    eyeId: eyeInput.value,
    screenCm: Number(screenInput.value),
    collimatorCm: Number(collimatorInput.value),
    objectCm: Number(objectInput.value),
    lensType: lensTypeInput.value,
    lensPower: Number(lensPowerInput.value),
    cylinderAngle: Number(cylinderInput.value)
  };
}

function updateExperiment(force = false) {
  const state = readExperimentState();
  const key = experimentKey(state);
  if (!force && key === lastExperimentKey) return;
  lastExperimentKey = key;

  setComponentPositions(state);
  updateLensControls();
  updateCorrectionSupport(state);
  updateSimulatedEyeSupport(state);

  const bundle = traceTeachingRays({
    eyeId: state.eyeId,
    lensType: state.lensType,
    lensPower: state.lensPower,
    screenCm: state.screenCm,
    cylinderAngle: state.cylinderAngle,
    objectCm: state.objectCm,
    collimatorCm: state.collimatorCm
  });

  updateRays(bundle);
  drawDetectorSpot(bundle.spot, bundle.result);
}
""",
    )

    add_code_block(
        doc,
        "代码片段 4：GLB 模型加载与场景挂载",
        """
export const OPTICS_MODELS = Object.freeze({
  bench: new URL('./assets/models/optics/bench.glb', import.meta.url).href,
  sourceParallel: new URL('./assets/models/optics/source-parallel.glb', import.meta.url).href,
  imageScreen: new URL('./assets/models/optics/image-screen.glb', import.meta.url).href,
  convexLens: new URL('./assets/models/optics/lens-convex.glb', import.meta.url).href,
  concaveLens: new URL('./assets/models/optics/lens-concave.glb', import.meta.url).href,
  cylinderLens: new URL('./assets/models/optics/lens-cylinder.glb', import.meta.url).href,
  eyeA: new URL('./assets/models/optics/sim-eye-a.glb', import.meta.url).href,
  eyeD: new URL('./assets/models/optics/sim-eye-d.glb', import.meta.url).href,
  eyeG: new URL('./assets/models/optics/sim-eye-g.glb', import.meta.url).href
});

export function attachOpticsModel(target, modelKey, options = {}) {
  target.userData.pendingModelKey = modelKey;
  loadOpticsModel(modelKey).then((model) => {
    if (!model || target.userData.pendingModelKey !== modelKey) return;
    applyModelTransform(model, options);
    target.add(model);
    if (options.hideFallback !== false) {
      target.children.forEach((child) => {
        if (child !== model && !child.userData.keepWithModel) child.visible = false;
      });
    }
    target.userData.loadedModel = model;
  });
}
""",
    )

    doc.add_heading("3.3 仿真平台操作流程", level=1)
    doc.add_paragraph(
        "平台操作流程对应真实实验教学过程：课前建立认知，实验中完成参数调节与现象观察，实验后进行数据处理和结论表达。"
    )
    for item in [
        "进入课前引导页面，了解实验目标和操作任务。",
        "进入基础知识页面，观察眼球结构及各部分在成像中的作用。",
        "进入模拟实验页面，选择 A-G 或 S 型模拟眼，调节像屏、物屏、准直透镜和矫正镜片。",
        "观察实时光路、焦点位置、光斑大小和读数面板，判断屈光状态。",
        "记录三次焦距测量值，系统自动计算平均焦距、焦度和推荐矫正值。",
        "进入实验报告页面，完成现象观察、数据分析、结论与反思。",
    ]:
        add_bullet(doc, item)

    add_figure(doc, "eye-layout-final.png", "图 3-5 基础知识页面：眼球结构与屈光系统展示")
    add_figure(doc, "ui-report-editable-tall.png", "图 3-6 实验报告页面：数据表与可编辑报告内容")

    doc.add_heading("3.4 基础性实验：模拟眼屈光状态观察", level=1)
    doc.add_paragraph(
        "基础性实验用于观察不同模拟眼的清晰成像位置。D 型模拟眼焦点与视网膜位置基本重合，"
        "A、B、C 型模拟眼焦点位于视网膜前，对应近视；E、F、G 型模拟眼焦点位于视网膜后，对应远视。"
        "学生通过移动像屏寻找清晰像位置，从现象层面理解屈光能力强弱与焦点位置的关系。"
    )

    doc.add_heading("3.5 进阶性实验：屈光不正矫正仿真", level=1)
    doc.add_paragraph(
        "进阶性实验用于研究不同屈光不正的矫正方法。近视眼使用凹透镜使光线适当发散，"
        "远视眼使用凸透镜使光线预先会聚，散光眼通过柱面镜及其轴向角度调节模拟不同方向屈光能力不一致的情况。"
    )

    doc.add_heading("3.6 数据处理与报告生成", level=1)
    doc.add_paragraph(
        "数据处理部分以焦距测量和焦度计算为核心。每个模拟眼可记录三次焦距测量值，系统自动计算平均值、"
        "焦度、屈光不正类型和矫正镜片焦度。报告页读取本地保存的数据，并提供可编辑文本框和打印导出功能。"
    )
    add_code_block(
        doc,
        "代码片段 5：实验数据记录与行计算",
        """
function applyRowCalculations(row, fittedPower = row.correctionFit) {
  if (!row.measurements.length) {
    resetRow(row);
    return row;
  }

  const average = row.measurements.reduce((sum, item) => sum + item, 0) / row.measurements.length;
  const diopter = diopterFromCm(average);
  const correction = correctionPowerForEye(average);

  row.average = average.toFixed(2);
  row.diopter = diopter.toFixed(2);
  row.type = classifyEye(average);
  row.correctionCalc = correction.toFixed(2);

  const fitted = Number(fittedPower);
  row.correctionFit = Number.isFinite(fitted) ? fitted.toFixed(2) : '';
  return row;
}
""",
    )
    add_code_block(
        doc,
        "代码片段 6：报告草稿保存与打印前同步",
        """
const DRAFT_KEY = 'eye-lab-report-draft-v1';

function loadDraft() {
  try {
    return JSON.parse(localStorage.getItem(DRAFT_KEY)) || {};
  } catch {
    return {};
  }
}

function saveDraft(draft) {
  localStorage.setItem(DRAFT_KEY, JSON.stringify(draft));
}

document.querySelectorAll('[data-report-draft]').forEach((field) => {
  const key = field.dataset.reportDraft;
  if (draft[key]) field.value = draft[key];
  field.addEventListener('input', () => {
    draft[key] = field.value;
    saveDraft(draft);
  });
});
""",
    )

    doc.add_heading("3.7 作品特色与创新点", level=1)
    add_bullet(doc, "三维可视化：通过 Three.js 呈现光具座、镜片、模拟眼和实时光路，增强实验直观性。")
    add_bullet(doc, "参数化交互：学生可以拖动器材或输入精确数值，实时观察光路和清晰度变化。")
    add_bullet(doc, "模型与数据联动：焦距、焦度、屈光类型和矫正焦度自动计算，降低机械计算负担。")
    add_bullet(doc, "报告一体化：实验数据、现象观察、分析结论和思考题集中到同一报告页面，支持完整实验流程。")

    doc.add_paragraph(
        "总体来看，本作品既复现了传统模拟眼实验的基本流程，又通过数字化方式扩展了实验的可操作性、"
        "可重复性和可视化程度，有助于学生理解屈光不正的物理本质，提高实验探究和数据分析能力。"
    )

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

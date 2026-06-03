from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "研究报告" / "实验数据处理与不确定度公式.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_equation(doc, formula):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(3)
    paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(formula)
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(13)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
    return paragraph


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(31, 77, 120)
    return paragraph


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("实验数据处理与不确定度公式")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(11, 37, 69)

    add_body(
        doc,
        "本页整理光学仿真实验报告中“实验数据处理与不确定度分析”部分可直接使用的核心公式。"
        "其中 f 表示焦距，D 表示屈光度，u 表示标准不确定度。"
    )

    add_heading(doc, "1. 多次测量平均焦距")
    add_body(doc, "对同一模拟眼重复测量三次焦点位置，取算术平均值作为实验焦距。")
    add_equation(doc, "f̄ = (f₁ + f₂ + f₃) / 3")

    add_heading(doc, "2. 焦距与屈光度换算")
    add_body(doc, "当焦距单位为 cm 时，屈光度可按下式换算：")
    add_equation(doc, "D = 100 / f̄")

    add_heading(doc, "3. 矫正镜片屈光度")
    add_body(doc, "以视网膜参考位置 f₀ = 24.0 cm 为目标，矫正镜片屈光度为：")
    add_equation(doc, "D_c = 100 / f₀ - 100 / f̄")
    add_body(doc, "若 D_c < 0，选择凹透镜；若 D_c > 0，选择凸透镜；若 D_c 接近 0，则无需明显矫正。")

    add_heading(doc, "4. A 类标准不确定度")
    add_body(doc, "重复测量数据的标准偏差为：")
    add_equation(doc, "s = √[ Σ(fᵢ - f̄)² / (n - 1) ]")
    add_body(doc, "平均焦距的 A 类标准不确定度为：")
    add_equation(doc, "u_A = s / √n")

    add_heading(doc, "5. B 类标准不确定度")
    add_body(doc, "若读数分辨率为 Δ，且误差服从均匀分布，则：")
    add_equation(doc, "u_B = Δ / √12")

    add_heading(doc, "6. 综合标准不确定度")
    add_body(doc, "焦距测量结果的综合标准不确定度为：")
    add_equation(doc, "u_f = √(u_A² + u_B²)")

    add_heading(doc, "7. 屈光度不确定度传递")
    add_body(doc, "由 D = 100 / f 可得屈光度的不确定度：")
    add_equation(doc, "u_D = |dD / df| · u_f = (100 / f²) · u_f")

    add_heading(doc, "8. 最终结果表达")
    add_body(doc, "实验结果可按“测量值 ± 标准不确定度”的形式表示：")
    add_equation(doc, "f = f̄ ± u_f")
    add_equation(doc, "D = D̄ ± u_D")

    add_heading(doc, "符号说明")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(3.0)
    table.columns[1].width = Cm(11.0)
    hdr = table.rows[0].cells
    hdr[0].text = "符号"
    hdr[1].text = "含义"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        set_cell_margin(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    rows = [
        ("f₁、f₂、f₃", "三次焦距测量值"),
        ("f̄", "平均焦距"),
        ("f₀", "视网膜参考位置，本系统取 24.0 cm"),
        ("D", "模拟眼屈光度"),
        ("D_c", "矫正镜片屈光度"),
        ("s", "重复测量标准偏差"),
        ("u_A", "A 类标准不确定度"),
        ("u_B", "B 类标准不确定度"),
        ("u_f", "焦距综合标准不确定度"),
        ("u_D", "屈光度标准不确定度"),
        ("Δ", "界面或刻度读数分辨率"),
        ("n", "重复测量次数"),
    ]
    for symbol, meaning in rows:
        cells = table.add_row().cells
        cells[0].text = symbol
        cells[1].text = meaning
        for cell in cells:
            set_cell_margin(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
